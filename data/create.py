import pandas as pd
import os
import glob
import re
import shutil
import numpy as np
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# --- KONFIGURASI ---
INPUT_FOLDER = 'Indonesian News Corpus/json'  # Folder data json
OUTPUT_FOLDER = 'dataset_ir_all' # Folder output baru

# Limit per kategori agar seimbang
# Total file nanti = Jumlah Kategori x TARGET_PER_CATEGORY
TARGET_PER_CATEGORY = 60
MIN_WORD_COUNT = 100      # Filter artikel pendek

def clean_text(text):
    """Membersihkan teks"""
    if not isinstance(text, str): return ""
    return text.replace('\r', '').strip()

def sanitize_filename(text):
    """Nama file aman"""
    clean = re.sub(r'[\\/*?:"<>|]', "", str(text))
    return clean[:80].strip()

def load_all_data(folder_path):
    print(f"1. Membaca seluruh data dari '{folder_path}'...")
    all_files = glob.glob(os.path.join(folder_path, "export-json-*"))
    
    if not all_files:
        print("   [ERROR] File JSON tidak ditemukan.")
        return pd.DataFrame()

    df_list = []
    for file in all_files:
        try:
            d = pd.read_json(file)
            df_list.append(d)
        except ValueError:
            try:
                d = pd.read_json(file, lines=True)
                df_list.append(d)
            except: pass

    if not df_list: return pd.DataFrame()

    full_df = pd.concat(df_list, ignore_index=True)
    
    # Preprocessing
    full_df.columns = full_df.columns.str.lower()
    if 'kategori' in full_df.columns:
        full_df['kategori'] = full_df['kategori'].astype(str).str.strip().str.title()
    
    # Hitung kata & Filter
    full_df['hitung_kata'] = full_df['isi'].apply(lambda x: len(str(x).split()))
    
    print(f"   Total data mentah: {len(full_df)}")
    
    # Hapus Duplikat & Filter Pendek
    full_df = full_df.drop_duplicates(subset=['isi'])
    full_df = full_df[full_df['hitung_kata'] >= MIN_WORD_COUNT]
    
    print(f"   Total data bersih (> {MIN_WORD_COUNT} kata): {len(full_df)}")
    return full_df

# --- FUNGSI SAVE ---
def save_txt(row, folder):
    fname = sanitize_filename(row['judul'])
    path = os.path.join(folder, f"{fname}.txt")
    try:
        with open(path, 'w', encoding='utf-8') as f:
            f.write(clean_text(row['isi']))
        return True
    except: return False

def save_docx(row, folder):
    fname = sanitize_filename(row['judul'])
    path = os.path.join(folder, f"{fname}.docx")
    try:
        doc = Document()
        doc.add_heading(row['judul'], level=1)
        doc.add_paragraph(f"Kategori: {row['kategori']} | Sumber: {row['sumber']}")
        doc.add_paragraph(clean_text(row['isi']))
        doc.save(path)
        return True
    except: return False

def save_pdf(row, folder):
    fname = sanitize_filename(row['judul'])
    path = os.path.join(folder, f"{fname}.pdf")
    try:
        doc = SimpleDocTemplate(path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Judul & Meta
        story.append(Paragraph(row['judul'], styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"<b>Kategori:</b> {row['kategori']}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Isi
        txt = clean_text(row['isi']).replace('\n', '<br/>')
        story.append(Paragraph(txt, styles['BodyText']))
        
        doc.build(story)
        return True
    except: return False

def main():
    # 1. Reset Folder
    if os.path.exists(OUTPUT_FOLDER):
        shutil.rmtree(OUTPUT_FOLDER)
    
    dirs = {
        'txt': os.path.join(OUTPUT_FOLDER, 'txt'),
        'docx': os.path.join(OUTPUT_FOLDER, 'docx'),
        'pdf': os.path.join(OUTPUT_FOLDER, 'pdf')
    }
    for d in dirs.values(): os.makedirs(d)

    # 2. Load Data
    df = load_all_data(INPUT_FOLDER)
    if df.empty: return

    # 3. Detect Categories
    categories = df['kategori'].unique()
    print(f"\n2. Ditemukan {len(categories)} Kategori: {categories}")
    print(f"   Target: {TARGET_PER_CATEGORY} file per kategori (dibagi 3 format).")

    total_files = 0

    # 4. Loop Semua Kategori
    for cat in categories:
        subset = df[df['kategori'] == cat]
        
        # Sampling
        if len(subset) < TARGET_PER_CATEGORY:
            print(f"   -> '{cat}': Data sedikit ({len(subset)}). Ambil semua.")
            samples = subset
        else:
            samples = subset.sample(n=TARGET_PER_CATEGORY, random_state=42)
        
        # Split ke 3 format
        parts = np.array_split(samples, 3)
        
        count_cat = 0
        # Batch 1 -> TXT
        for _, row in parts[0].iterrows():
            if save_txt(row, dirs['txt']): count_cat += 1
        # Batch 2 -> DOCX
        for _, row in parts[1].iterrows():
            if save_docx(row, dirs['docx']): count_cat += 1
        # Batch 3 -> PDF
        for _, row in parts[2].iterrows():
            if save_pdf(row, dirs['pdf']): count_cat += 1
            
        print(f"      Disimpan: {count_cat} file.")
        total_files += count_cat

    print("-" * 40)
    print(f"SELESAI! Total {total_files} file tersimpan di '{OUTPUT_FOLDER}'.")

if __name__ == "__main__":
    main()